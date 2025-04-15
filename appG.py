import os
import tempfile
import requests
import base64
import mimetypes
from flask import Flask, request, jsonify
from dotenv import load_dotenv
from pdf2image import convert_from_path
from PIL import Image
import pandas as pd
import json
from pytesseract import image_to_string
import pytesseract
from urllib.parse import unquote, urlparse
from google import genai
from google.genai import types
import google.generativeai as genai
from PIL import Image
from docx import Document

app = Flask(__name__)
load_dotenv()

# Initialize Gemini Client
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

EXCEL_PATH = "D:\\Bikham\\Map_Proj\\Doc Types (3).xlsx"  # Replace with your local Excel file path
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

SYSTEM_PROMPT = """
You are a Senior Healthcare Document Specialist with 20+ years experience. 
You must STRICTLY match documents to one of the following standard names:

<STANDARD NAMES>
{standard_names}
</STANDARD NAMES>

### Absolute Rules:
1. 🔒 Only return names EXACTLY as written above — no shortening or rewording.
2. 🔍 Match key identifiers:
   - Form numbers (CMS-XXXX, I-XXX)
   - Government agencies (CMS, Medicare, Medicaid, DEA, IRS)
   - Document type (Certification, Enrollment, Claim Form, etc.)
3. 🥇 Choose the most *specific* match available from the list.

### Match Procedure:
- Review full document structure and content
- Compare headers, titles, and context to the standard list
- Pick the entry that aligns **word-for-word**

### Output Format (JSON only):
```json
{{
  "document_name": "MUST MATCH EXACTLY from standard list",
  "match_reason": "Brief reason (25 words max)"
}}

Note: 
Use the sent Document name to identify the standard document name. Sent name can consist provider name in start so ignore that.
"""

def extract_filename_from_firebase_url(url):
    path = urlparse(url).path
    decoded_path = unquote(path)
    filename = decoded_path.split('/')[-1]
    return filename

def extract_text_with_ocr(image_path):
    try:
        with Image.open(image_path) as img:
            img = img.convert('L')
            img = img.point(lambda x: 0 if x < 128 else 255)
            extracted_text = pytesseract.image_to_string(
                img,
                config='--psm 3 --oem 3 -c preserve_interword_spaces=1'
            )
            return extracted_text.strip()
    except Exception as e:
        raise RuntimeError(f"OCR failed: {str(e)}")

def download_file_from_url(url):
    response = requests.get(url, stream=True)
    if response.status_code != 200:
        raise Exception("Failed to download the file from the provided URL.")
    content_type = response.headers.get("Content-Type", "")
    print("content type : ", content_type)
    ext = mimetypes.guess_extension(content_type.split(";")[0]) or ".pdf"
    print("ext : ", ext)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    with open(temp_file.name, 'wb') as f:
        f.write(response.content)
    return temp_file.name

def convert_pdf_to_image(pdf_path):
    images = convert_from_path(pdf_path, first_page=1, last_page=1)
    temp_img = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    images[0].save(temp_img.name, "PNG")
    return temp_img.name

def get_image_path(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type == "application/pdf":
        return convert_pdf_to_image(file_path)
    elif mime_type and mime_type.startswith("image"):
        return file_path
    else:
        raise ValueError("Unsupported file format")

def get_standard_names():
    df = pd.read_excel(EXCEL_PATH)
    return df["name"].astype(str).tolist()

def regex_search(var):
    text = str(var)
    start_index = text.find('{')
    end_index = text.rfind('}')
    clean_text = text[start_index:end_index + 1]
    data = json.loads(clean_text)
    return data

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".pdf", ".jpg", ".jpeg", ".png"]:
        image_path = get_image_path(file_path)
        return extract_text_with_ocr(image_path), Image.open(image_path)

    elif ext == ".csv":
        df = pd.read_csv(file_path)
        return df.to_string(index=False), None

    elif ext == ".xlsx":
        df = pd.read_excel(file_path)
        return df.to_string(index=False), None

    elif ext == ".docx":
        doc = Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        return full_text, None

    else:
        raise ValueError("Unsupported file format")

def identify_document_name(file_path, pdfName):
    standard_names = get_standard_names()
    formatted_names = "\n".join(f'"{name}"' for name in standard_names)
    prompt = SYSTEM_PROMPT.replace("{standard_names}", formatted_names)

    try:
        ocr_text, image_obj = extract_text_from_file(file_path)
        print("ocr text : ", ocr_text)

        model = genai.GenerativeModel("gemini-2.5-pro-preview-03-25")

        inputs = [
            prompt,
            f"OCR EXTRACT:\n{ocr_text}",
            f"Sent Document(pdf/jpg/png/csv/xlsx/docx):\n{pdfName}",
        ]
        if image_obj:
            inputs.append(image_obj)

        response = model.generate_content(inputs)

        return json.dumps(regex_search(response.text))

    except Exception as e:
        return json.dumps({
            "document_name": "UNKNOWN",
            "notes": str(e)
        })

@app.route("/identify", methods=["POST"])
def identify():
    data = request.json
    file_url = data.get("file_url")
    local_path = None

    if not file_url:
        return jsonify({"error": "Missing 'file_url' in request."}), 400

    try:
        pdfName = extract_filename_from_firebase_url(file_url)
        print("Sent Document Name : ", pdfName)
        local_path = download_file_from_url(file_url)
        matched_name = identify_document_name(local_path, pdfName)
        return jsonify(json.loads(matched_name))
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if local_path and os.path.exists(local_path):
            os.remove(local_path)

if __name__ == "__main__":
    app.run(debug=True)
